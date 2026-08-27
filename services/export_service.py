"""
export_service.py
------------------
Generates downloadable, properly formatted meeting-notes documents:
PDF (reportlab), DOCX (python-docx), and plain TXT.
"""

import os

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, Table, TableStyle
)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PURPLE = "6C5CE7"


def _speakers_text(meeting):
    return ", ".join(s["name"] for s in meeting.get("speakers", [])) or "Not identified"


def build_text_notes(meeting):
    """Returns a plain-text string with the full formatted meeting notes."""
    lines = []
    lines.append(f"MEETING NOTES - {meeting.get('title', 'Untitled Meeting')}")
    lines.append(f"Date & Time: {meeting.get('date', '')}")
    lines.append(f"Duration: {meeting.get('duration', 0)} seconds")
    lines.append(f"Participants: {_speakers_text(meeting)}")
    lines.append("")
    lines.append("SUMMARY")
    lines.append(meeting.get("summary") or "No summary available.")
    lines.append("")
    lines.append("ACTION ITEMS")
    items = meeting.get("action_items") or []
    if items:
        for i in items:
            box = "[x]" if i.get("completed") else "[ ]"
            lines.append(f"  {box} {i.get('task','')} - {i.get('assigned_to','Unassigned')} (Due: {i.get('deadline','N/A')})")
    else:
        lines.append("  None recorded.")
    lines.append("")
    lines.append("KEY HIGHLIGHTS")
    highlights = meeting.get("key_highlights") or []
    if highlights:
        for h in highlights:
            lines.append(f"  - {h}")
    else:
        lines.append("  None recorded.")
    lines.append("")
    lines.append("DECISIONS")
    decisions = meeting.get("decisions") or []
    if decisions:
        for d in decisions:
            lines.append(f"  - {d}")
    else:
        lines.append("  None recorded.")
    lines.append("")
    lines.append("FULL TRANSCRIPT")
    lines.append(meeting.get("transcript") or "No transcript available.")
    return "\n".join(lines)


def export_txt(meeting, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(build_text_notes(meeting))
    return out_path


def export_pdf(meeting, out_path):
    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm, leftMargin=2 * cm, rightMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle", parent=styles["Title"], textColor=colors.HexColor(f"#{PURPLE}"), fontSize=22,
    )
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=colors.HexColor(f"#{PURPLE}"), spaceBefore=14)
    body = styles["BodyText"]
    body.leading = 15

    story = [
        Paragraph(meeting.get("title", "Untitled Meeting"), title_style),
        Spacer(1, 6),
        Paragraph(f"<b>Date &amp; Time:</b> {meeting.get('date','')}", body),
        Paragraph(f"<b>Duration:</b> {meeting.get('duration',0)} seconds", body),
        Paragraph(f"<b>Participants:</b> {_speakers_text(meeting)}", body),
    ]

    story.append(Paragraph("Summary", h2))
    story.append(Paragraph(meeting.get("summary") or "No summary available.", body))

    story.append(Paragraph("Action Items", h2))
    items = meeting.get("action_items") or []
    if items:
        rows = [["Done", "Task", "Assigned To", "Deadline"]]
        for i in items:
            rows.append(["Yes" if i.get("completed") else "No", i.get("task", ""), i.get("assigned_to", ""), i.get("deadline", "")])
        table = Table(rows, colWidths=[2 * cm, 7 * cm, 4 * cm, 3 * cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{PURPLE}")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E0E0E0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
    else:
        story.append(Paragraph("None recorded.", body))

    story.append(Paragraph("Key Highlights", h2))
    highlights = meeting.get("key_highlights") or []
    if highlights:
        story.append(ListFlowable([ListItem(Paragraph(h, body)) for h in highlights], bulletType="bullet"))
    else:
        story.append(Paragraph("None recorded.", body))

    story.append(Paragraph("Decisions Made", h2))
    decisions = meeting.get("decisions") or []
    if decisions:
        story.append(ListFlowable([ListItem(Paragraph(d, body)) for d in decisions], bulletType="bullet"))
    else:
        story.append(Paragraph("None recorded.", body))

    story.append(Paragraph("Full Transcript", h2))
    transcript = (meeting.get("transcript") or "No transcript available.").replace("\n", "<br/>")
    story.append(Paragraph(transcript, body))

    doc.build(story)
    return out_path


def export_docx(meeting, out_path):
    doc = Document()

    title = doc.add_heading(meeting.get("title", "Untitled Meeting"), level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x6C, 0x5C, 0xE7)

    meta = doc.add_paragraph()
    meta.add_run(f"Date & Time: {meeting.get('date','')}\n").bold = True
    meta.add_run(f"Duration: {meeting.get('duration',0)} seconds\n").bold = True
    meta.add_run(f"Participants: {_speakers_text(meeting)}").bold = True

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(meeting.get("summary") or "No summary available.")

    doc.add_heading("Action Items", level=1)
    items = meeting.get("action_items") or []
    if items:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 5"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Done", "Task", "Assigned To", "Deadline"
        for i in items:
            row = table.add_row().cells
            row[0].text = "Yes" if i.get("completed") else "No"
            row[1].text = i.get("task", "")
            row[2].text = i.get("assigned_to", "")
            row[3].text = i.get("deadline", "")
    else:
        doc.add_paragraph("None recorded.")

    doc.add_heading("Key Highlights", level=1)
    highlights = meeting.get("key_highlights") or []
    if highlights:
        for h in highlights:
            doc.add_paragraph(h, style="List Bullet")
    else:
        doc.add_paragraph("None recorded.")

    doc.add_heading("Decisions Made", level=1)
    decisions = meeting.get("decisions") or []
    if decisions:
        for d in decisions:
            doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("None recorded.")

    doc.add_heading("Full Transcript", level=1)
    doc.add_paragraph(meeting.get("transcript") or "No transcript available.")

    doc.save(out_path)
    return out_path
